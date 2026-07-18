from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont


OUT = "/Users/master/CodeBuddy/art12/100艺术官方网站介绍文档-初稿.docx"
LOGO = "/Users/master/CodeBuddy/art12/frontend/src/static/landing/yiben-logo.jpg"
FLOWCHART = "/Users/master/CodeBuddy/art12/100艺术-作品流动与收益流程图.png"

INK = "17140F"
GOLD = "B88A32"
DEEP_GOLD = "7A5717"
CREAM = "F6F0E3"
PALE = "FBF8F1"
GRAY = "6F6A61"
LINE = "D9CEB8"
WHITE = "FFFFFF"


def generate_flowchart(path):
    width, height = 1800, 1180
    image = Image.new("RGB", (width, height), "#FCFAF5")
    draw = ImageDraw.Draw(image)
    font_path = "/Library/Fonts/Arial Unicode.ttf"
    title_font = ImageFont.truetype(font_path, 52)
    label_font = ImageFont.truetype(font_path, 31)
    body_font = ImageFont.truetype(font_path, 25)
    small_font = ImageFont.truetype(font_path, 22)
    bold_font = ImageFont.truetype(font_path, 30)

    draw.text((90, 50), "作品流动与多角色收益闭环", font=title_font, fill="#17140F")
    draw.text((92, 116), "作品每次成交都会形成新的认证、价格与流通记录，并产生对应收益。",
              font=body_font, fill="#6F6A61")

    process = [
        ("艺术家发布作品", 80),
        ("首次成交", 370),
        ("收藏家持有", 660),
        ("申请再次出售", 950),
        ("艺术家认证", 1240),
        ("再次成交", 1530),
    ]
    box_y, box_w, box_h = 205, 210, 94
    for idx, (text, x) in enumerate(process):
        draw.rounded_rectangle((x, box_y, x + box_w, box_y + box_h), radius=20,
                               fill="#F2EEE5", outline="#9A9388", width=3)
        bbox = draw.textbbox((0, 0), text, font=label_font)
        draw.text((x + (box_w - (bbox[2] - bbox[0])) / 2,
                   box_y + (box_h - (bbox[3] - bbox[1])) / 2 - 4),
                  text, font=label_font, fill="#27231D")
        if idx < len(process) - 1:
            ax1, ax2, ay = x + box_w + 18, process[idx + 1][1] - 18, box_y + box_h / 2
            draw.line((ax1, ay, ax2, ay), fill="#A8791D", width=5)
            draw.polygon([(ax2, ay), (ax2 - 17, ay - 11), (ax2 - 17, ay + 11)], fill="#A8791D")

    draw.text((90, 350), "收益分配", font=bold_font, fill="#17140F")

    lane_specs = [
        ("艺术家", "#D39A29", "#FFF2CC", [
            ("收益点 ①  首次销售收入", "获得作品首次成交的销售款"),
            ("收益点 ②  再次流通认证费", "参考：再次交易金额的 2%–5%"),
            ("长期价值积累", "认证、价格与成交记录持续沉淀"),
        ]),
        ("经纪人", "#4E78AD", "#E7F0FC", [
            ("首次成交分销费", "参考：成交金额的 5%–15%"),
            ("再次成交分销费", "每次促成交易均可获得服务回报"),
        ]),
        ("收藏家", "#5E8962", "#E6F2E5", [
            ("获得并持有作品", "拥有收藏、展示和再次流通权利"),
            ("藏品再次出售收入", "转售成交后获得藏品销售款"),
        ]),
    ]
    lane_y = 410
    lane_h = 205
    for lane_idx, (role, color, pale, items) in enumerate(lane_specs):
        y = lane_y + lane_idx * 235
        draw.rounded_rectangle((80, y, 1720, y + lane_h), radius=24,
                               fill="#FFFFFF", outline=color, width=4)
        draw.rounded_rectangle((80, y, 290, y + lane_h), radius=24, fill=color)
        draw.rectangle((260, y, 310, y + lane_h), fill=color)
        rb = draw.textbbox((0, 0), role, font=bold_font)
        draw.text((185 - (rb[2] - rb[0]) / 2, y + 76), role, font=bold_font, fill="#FFFFFF")

        count = len(items)
        start_x, available = 330, 1360
        gap = 24
        item_w = (available - gap * (count - 1)) / count
        for i, (heading, detail) in enumerate(items):
            x = start_x + i * (item_w + gap)
            draw.rounded_rectangle((x, y + 28, x + item_w, y + lane_h - 28), radius=18,
                                   fill=pale, outline=color, width=2)
            hb = draw.textbbox((0, 0), heading, font=body_font)
            draw.text((x + (item_w - (hb[2] - hb[0])) / 2, y + 58), heading,
                      font=body_font, fill="#17140F")
            db = draw.textbbox((0, 0), detail, font=small_font)
            draw.text((x + (item_w - (db[2] - db[0])) / 2, y + 112), detail,
                      font=small_font, fill="#514B43")

    draw.text((90, 1120), "说明：费用比例为当前方案的参考区间，最终以平台正式规则和具体服务内容为准。",
              font=small_font, fill="#6F6A61")
    image.save(path, quality=95)


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=130, start=160, bottom=130, end=160):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_font(run, size=11, bold=False, color=INK, name="STSong"):
    run.font.name = name
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), name)
    run._element.rPr.rFonts.set(qn("w:ascii"), name)
    run._element.rPr.rFonts.set(qn("w:hAnsi"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)


def add_text(doc, text, size=11, bold=False, color=INK, after=8, before=0,
             align=WD_ALIGN_PARAGRAPH.LEFT, line=1.35, keep=False):
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = line
    p.paragraph_format.keep_with_next = keep
    set_font(p.add_run(text), size=size, bold=bold, color=color)
    return p


def add_heading(doc, text, level=1):
    size = {1: 18, 2: 14, 3: 11.5}[level]
    before = {1: 18, 2: 13, 3: 9}[level]
    after = {1: 8, 2: 6, 3: 4}[level]
    p = add_text(doc, text, size=size, bold=True,
                 color=DEEP_GOLD if level < 3 else INK,
                 before=before, after=after, line=1.1, keep=True)
    p.style = f"Heading {level}"
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.28)
    p.paragraph_format.first_line_indent = Inches(-0.18)
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(text), size=10.5)
    return p


def add_callout(doc, label, headline, body):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)
    cell = table.cell(0, 0)
    cell.width = Inches(6.5)
    set_cell_shading(cell, CREAM)
    set_cell_margins(cell, 180, 240, 180, 240)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(label.upper()), size=9, bold=True, color=GOLD)
    p = cell.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run(headline), size=16, bold=True, color=INK)
    p = cell.add_paragraph()
    p.paragraph_format.line_spacing = 1.3
    set_font(p.add_run(body), size=10.5, color=GRAY)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_mechanism(doc, number, title, tagline, details, revenue):
    table = doc.add_table(rows=1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [0.72, 5.78]
    for i, width in enumerate(widths):
        table.columns[i].width = Inches(width)
        table.cell(0, i).width = Inches(width)
        set_cell_margins(table.cell(0, i), 170, 180, 170, 180)
    set_cell_shading(table.cell(0, 0), GOLD)
    set_cell_shading(table.cell(0, 1), PALE)
    p = table.cell(0, 0).paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(f"{number:02d}"), size=16, bold=True, color=WHITE)
    c = table.cell(0, 1)
    p = c.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    set_font(p.add_run(title), size=13, bold=True, color=DEEP_GOLD)
    p = c.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    set_font(p.add_run(tagline), size=10.5, bold=True, color=INK)
    p = c.add_paragraph()
    p.paragraph_format.space_after = Pt(5)
    p.paragraph_format.line_spacing = 1.25
    set_font(p.add_run(details), size=10, color=GRAY)
    p = c.add_paragraph()
    set_font(p.add_run(revenue), size=9.5, bold=True, color=DEEP_GOLD)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


generate_flowchart(FLOWCHART)
doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.78)
section.bottom_margin = Inches(0.72)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.35)
section.footer_distance = Inches(0.35)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "STSong"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "STSong")
normal.font.size = Pt(11)
normal.font.color.rgb = RGBColor.from_string(INK)
for name in ("Heading 1", "Heading 2", "Heading 3"):
    style = styles[name]
    style.font.name = "STSong"
    style._element.rPr.rFonts.set(qn("w:eastAsia"), "STSong")

header = section.header.paragraphs[0]
header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
set_font(header.add_run("100艺术｜官方网站介绍文档"), size=8.5, color=GRAY)
footer = section.footer.paragraphs[0]
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_font(footer.add_run("100艺术  ·  让创作永远有价值"), size=8.5, color=GRAY)

# Cover
add_text(doc, "BRAND INTRODUCTION", size=9, bold=True, color=GOLD, before=18,
         after=28, align=WD_ALIGN_PARAGRAPH.CENTER)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run()
r.add_picture(LOGO, width=Inches(3.6))
p.paragraph_format.space_after = Pt(26)
add_text(doc, "100艺术官方网站介绍", size=28, bold=True, color=INK, after=9,
         align=WD_ALIGN_PARAGRAPH.CENTER, line=1.05)
add_text(doc, "让创作永远有价值", size=17, bold=True, color=DEEP_GOLD, after=18,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_text(doc, "连接艺术家、经纪人和收藏家，建立作品从首次成交、真伪认证到再次流通的长期价值网络。",
         size=11.5, color=GRAY, after=36, align=WD_ALIGN_PARAGRAPH.CENTER, line=1.5)
add_callout(doc, "核心主张", "艺术作品不止完成一次交易",
            "100艺术希望让每一次创作都拥有可验证的身份、可持续的传播、可更新的价值与可再次流通的机会。")
add_text(doc, "官方介绍文档 · 第一版 · 2026年7月", size=9, color=GRAY, before=28,
         align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_page_break()

add_heading(doc, "一、100艺术是什么", 1)
add_text(doc, "100艺术是面向艺术家、艺术经纪人与收藏家的艺术品展示、交易、认证和再流通平台。平台不把作品视为一次性售卖的商品，而是围绕作品的长期身份、市场传播与流通记录，构建可持续的艺术价值体系。")
add_text(doc, "在这个体系中，艺术家可以持续参与作品的生命旅程；经纪人可以凭借专业传播能力获得合理回报；收藏家可以更安心地收藏，也可以让已收藏作品重新进入市场。")
add_callout(doc, "一句话定位", "让每一件作品拥有持续生长的价值",
            "创作被看见，交易可追溯，真伪可认证，价格有依据，收藏可流通。")

add_heading(doc, "二、品牌核心：什么是“永远有价值”", 1)
add_text(doc, "“永远有价值”并不意味着平台承诺作品必然升值，而是指艺术家完成创作和首次销售之后，仍然能够通过认证、传播、市场影响力和作品再次流通，持续参与作品价值的形成。")
add_bullet(doc, "作品有身份：建立艺术家、作品信息、认证记录与流通记录之间的关联。")
add_bullet(doc, "艺术家有持续收益：作品再次流通且需要艺术家进行真伪认证时，艺术家可获得认证服务费用。")
add_bullet(doc, "价值有依据：平台根据可量化信息持续更新作品的市场参考价格。")
add_bullet(doc, "收藏有出口：收藏家可以将符合条件的藏品再次上架，进入平台流通体系。")

add_heading(doc, "三、平台四大核心机制", 1)
add_mechanism(doc, 1, "艺术家持续认证收益", "作品售出以后，艺术家仍然与作品保持价值连接。",
              "当作品再次在平台流通，并需要原作者或授权主体确认真伪时，平台发起标准化认证流程。认证结果及相关记录可成为作品后续交易的重要凭证。",
              "参考机制：艺术家可获得该次交易金额约 2%–5% 的认证服务费用，具体比例以平台规则、作品类别及服务内容为准。")
add_mechanism(doc, 2, "经纪人分销协作", "让专业的人帮助艺术家被更多人看见。",
              "平台聚合艺术经纪人、推广者与合作伙伴。经纪人可以选择适合其客户和渠道的作品进行推广，并通过平台记录推广关系、成交归属和服务过程。",
              "参考机制：经纪人促成作品成交后，可获得约 5%–15% 的分销服务费用，具体比例由作品方设置或按平台合作规则执行。")
add_mechanism(doc, 3, "作品价格动态更新", "作品价格不是静止标签，而是艺术家发展轨迹的市场表达。",
              "平台综合艺术家的创作年限、公开展览与职业经历、互联网关注度、身份与荣誉变化、作品成交及流通记录等信息，形成动态市场参考价格。",
              "价格说明：平台展示的是辅助交易决策的参考信息，不构成保值、升值或投资收益承诺。")
add_mechanism(doc, 4, "收藏家藏品再流通", "收藏不再只有进入，也拥有规范、可信的退出通道。",
              "收藏家可将已收藏且符合平台规则的作品申请再次售卖。平台可协助完成信息核验、真伪认证、价格参考、展示推广和交易服务。",
              "流通价值：提升作品信息透明度与交易效率，让作品在艺术家、经纪人和收藏家之间持续流动。")

doc.add_page_break()

add_heading(doc, "四、平台价值闭环", 1)
add_text(doc, "一件作品从首次发布到再次流通，会在不同阶段为艺术家、经纪人和收藏家形成清晰的收益点。",
         size=10.5, color=GRAY, after=8)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(5)
r = p.add_run()
r.add_picture(FLOWCHART, width=Inches(6.5))
add_text(doc, "图 1｜作品流动与多角色收益闭环", size=9, color=GRAY, after=10,
         align=WD_ALIGN_PARAGRAPH.CENTER)
add_text(doc, "100艺术围绕一件作品建立完整旅程：")
steps = [
    ("01", "创作与确权", "艺术家发布作品，建立基础信息、作者关系与作品档案。"),
    ("02", "展示与分销", "作品通过平台和经纪人网络获得更广泛的展示与成交机会。"),
    ("03", "首次成交", "交易完成后沉淀作品、买家、价格和服务记录。"),
    ("04", "持续估值", "根据艺术家发展及市场信息更新作品参考价格。"),
    ("05", "再次流通", "收藏家申请转售，平台组织核验、展示和交易。"),
    ("06", "认证与收益", "需要真伪认证时，艺术家提供认证服务并获得相应费用。"),
]
table = doc.add_table(rows=1, cols=3)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.autofit = False
widths = [0.65, 1.35, 4.5]
for i, width in enumerate(widths):
    table.columns[i].width = Inches(width)
    table.cell(0, i).width = Inches(width)
for i, text in enumerate(("阶段", "节点", "平台作用")):
    cell = table.cell(0, i)
    set_cell_shading(cell, DEEP_GOLD)
    set_cell_margins(cell)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(p.add_run(text), size=10, bold=True, color=WHITE)
set_repeat_table_header(table.rows[0])
for n, title, body in steps:
    cells = table.add_row().cells
    for i, width in enumerate(widths):
        cells[i].width = Inches(width)
        set_cell_margins(cells[i])
        set_cell_shading(cells[i], PALE if int(n) % 2 else WHITE)
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_font(cells[0].paragraphs[0].add_run(n), size=10, bold=True, color=GOLD)
    set_font(cells[1].paragraphs[0].add_run(title), size=10, bold=True)
    set_font(cells[2].paragraphs[0].add_run(body), size=9.5, color=GRAY)

add_heading(doc, "五、平台为谁创造价值", 1)
add_heading(doc, "艺术家", 2)
add_text(doc, "获得长期作品档案、展示成交渠道、经纪人协作网络、市场价格参考，以及作品再次流通时的认证服务收益。", after=5)
add_heading(doc, "经纪人", 2)
add_text(doc, "获得可选择、可推广、可追踪的艺术作品资源，通过专业匹配、传播和成交服务获得分销回报。", after=5)
add_heading(doc, "收藏家", 2)
add_text(doc, "获得更透明的作品信息、更可信的认证路径、更清晰的价格参考，以及收藏作品再次流通的渠道。", after=5)

add_heading(doc, "六、官网建议内容架构", 1)
add_text(doc, "后续官方网站可按照“先讲价值，再讲机制，最后促成行动”的顺序组织内容：")
for item in [
    "首屏：100艺术｜让创作永远有价值。",
    "品牌解释：用一段话回答“为什么作品成交后仍然与艺术家有关”。",
    "四大机制：持续认证收益、经纪人分销、动态价格、藏品再流通。",
    "价值闭环：以作品旅程图展示从创作到再次流通的全过程。",
    "角色入口：艺术家入驻、经纪人合作、收藏家浏览与转售。",
    "信任说明：认证流程、价格方法、费用规则、交易记录和风险提示。",
    "行动入口：进入 App、艺术家入驻、经纪人合作、藏品送售。",
]:
    add_bullet(doc, item)

doc.add_page_break()

add_heading(doc, "七、官方网站标准介绍文案", 1)
add_callout(doc, "官网主标题", "让创作永远有价值",
            "100艺术连接艺术家、经纪人和收藏家，让作品从首次展示与成交，走向可信认证、动态定价和长期流通。")

add_heading(doc, "品牌长介绍", 2)
add_text(doc, "100艺术相信，一件作品的价值不应止于第一次成交。艺术家完成创作并售出作品后，当作品再次进入市场流通并需要真伪认证时，艺术家仍可通过提供认证服务获得合理回报；平台上的经纪人可以帮助艺术家推广和分销作品，并在促成成交后获得服务费用；作品的市场参考价格将结合艺术家的创作年限、职业经历、互联网关注度、身份变化及成交记录持续更新；收藏家也可以将符合条件的藏品再次带回平台，完成认证、展示与售卖。")
add_text(doc, "我们希望建立的，不只是一个艺术品交易平台，而是一套让创作被尊重、让专业被回报、让收藏可流通的长期价值体系。")

add_heading(doc, "品牌短介绍", 2)
add_text(doc, "100艺术是连接艺术家、经纪人与收藏家的艺术品长期流通平台。通过艺术家认证收益、经纪人分销、动态价格参考和收藏品再流通机制，让作品在每一次被看见、被收藏和再次交易时持续产生价值。")

add_heading(doc, "四句核心传播语", 2)
add_bullet(doc, "对艺术家：作品售出，不代表价值关系结束。")
add_bullet(doc, "对经纪人：让专业推广获得清晰、持续的回报。")
add_bullet(doc, "对收藏家：让每一件收藏都有可信记录与流通出口。")
add_bullet(doc, "对市场：让艺术价值有记录、有依据、可延续。")

add_heading(doc, "八、正式发布前需要确认的规则", 1)
add_text(doc, "为保证官网表述准确、透明，以下内容应在产品规则和合同确定后再对外发布：")
add_bullet(doc, "艺术家认证服务费的计算基数、适用条件、比例范围及结算方式。")
add_bullet(doc, "经纪人分销比例的设置权限、归因规则、退款及售后情况下的处理方式。")
add_bullet(doc, "动态参考价格的数据来源、更新周期、计算维度及人工复核机制。")
add_bullet(doc, "收藏品再次上架的准入条件、认证责任、交付方式和平台服务费用。")
add_bullet(doc, "作品真伪、价格信息、交易风险及非投资承诺等必要提示。")

add_callout(doc, "文案边界", "“永远有价值”是品牌使命，不是升值承诺",
            "官网应持续强调长期记录、认证服务和流通机会，避免使用“保证升值”“固定回报”“稳赚”等可能引发误解的表达。")

doc.save(OUT)
print(OUT)
